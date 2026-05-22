#!/usr/bin/env python3
"""
mdtool - manual CLI utility for working with markdown stored in a local
GitHub repository using Microsoft Word on Windows.

Supports:
    - configuration mode (--config, or auto-triggered when config.ini is absent)
    - --fromgit: bootstrap .docx working copies from .md files
    - --togit: convert edited .docx files back to .md, extracting and
      compositing images

Out of scope: git/GitHub integration, recursive directory traversal,
filesystem monitoring.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    from markdown_it import MarkdownIt
    from markdown_it.token import Token
except ImportError:
    sys.stderr.write(
        "ERROR: markdown-it-py is required.\n"
        "Install with: pip install markdown-it-py\n"
    )
    sys.exit(2)

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is required.\n"
        "Install with: pip install python-docx\n"
    )
    sys.exit(2)

try:
    from PIL import Image
except ImportError:
    sys.stderr.write(
        "ERROR: Pillow is required (for --togit image rendering).\n"
        "Install with: pip install Pillow\n"
    )
    sys.exit(2)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CONFIG_FILENAME = "config.ini"
CONFIG_KEYS: Tuple[str, ...] = ("GitDir", "WorkDir")

CODE_FONT = "Consolas"
CODE_BLOCK_STYLE_NAME = "CodeBlock"
INLINE_CODE_STYLE_NAME = "InlineCode"
TABLE_STYLE_NAME = "Table Grid"
QUOTE_STYLE_NAME = "Quote"

CODE_SHADING_FILL = "F2F2F2"
HR_BORDER_COLOR = "808080"

# --togit
ASSETS_DIR_NAME = "assets"
MANIFEST_FILENAME = ".mdwordtool-manifest.json"
MANIFEST_VERSION = 1
STATE_FILENAME = ".mdwordtool-state.json"
STATE_VERSION = 1
IMAGE_DPI = 150  # Render images at this DPI of Word display size.
SUPPORTED_IMAGE_CONTENT_TYPES = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
}
LIST_INDENT_STEP_CM = 0.75  # Matches --fromgit indent step.


# ---------------------------------------------------------------------------
# Error handling
# ---------------------------------------------------------------------------

def fail(message: str) -> None:
    """Print error to stderr and exit non-zero. Fail-fast per spec."""
    sys.stderr.write(f"ERROR: {message}\n")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Config handling
# ---------------------------------------------------------------------------

def config_path() -> Path:
    return Path.cwd() / CONFIG_FILENAME


def load_config(path: Path) -> Optional[Dict[str, str]]:
    """Load config from key=value INI-style file. Returns None if file absent."""
    if not path.exists():
        return None
    cfg: Dict[str, str] = {}
    try:
        with open(path, "r", encoding="utf-8") as fp:
            for raw in fp:
                line = raw.strip()
                if not line or line.startswith("#") or line.startswith(";"):
                    continue
                if "=" not in line:
                    continue
                k, _, v = line.partition("=")
                cfg[k.strip()] = v.strip()
    except OSError as exc:
        fail(f"Failed to read {path}: {exc}")
    return cfg


def write_config(path: Path, cfg: Dict[str, str]) -> None:
    try:
        with open(path, "w", encoding="utf-8") as fp:
            for key in CONFIG_KEYS:
                fp.write(f"{key} = {cfg[key]}\n")
    except OSError as exc:
        fail(f"Failed to write {path}: {exc}")


def run_config_mode(existing: Optional[Dict[str, str]]) -> None:
    print("Entering configuration mode.")
    new_cfg: Dict[str, str] = {}

    for key in CONFIG_KEYS:
        default = (existing or {}).get(key, "")
        prompt = key
        if default:
            prompt += f" [{default}]"
        prompt += ": "
        try:
            value = input(prompt).strip()
        except EOFError:
            fail("Input ended unexpectedly during configuration.")
            return  # unreachable
        if not value:
            value = default
        if not value:
            fail(f"{key} must not be empty.")
        new_cfg[key] = value

    print("\nProposed configuration:")
    for key in CONFIG_KEYS:
        print(f"  {key} = {new_cfg[key]}")

    try:
        confirm = input("Confirm and write config? [y/n]: ").strip().lower()
    except EOFError:
        fail("Input ended unexpectedly during confirmation.")
        return  # unreachable

    if confirm != "y":
        print("Configuration cancelled.")
        sys.exit(0)

    # Validate that all configured paths exist as directories.
    for key in CONFIG_KEYS:
        p = Path(new_cfg[key])
        if not p.is_dir():
            fail(
                f"Path for {key} does not exist or is not a directory: "
                f"{new_cfg[key]}"
            )

    write_config(config_path(), new_cfg)
    print(f"Configuration written to {config_path()}.")
    sys.exit(0)


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------

def _shading(fill_hex: str) -> "OxmlElement":
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    return shd


def _add_paragraph_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_pr.append(pbdr)
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), HR_BORDER_COLOR)
    pbdr.append(border)


def _mark_quick_format(style) -> None:
    """Add <w:qFormat/> so the style appears in Word's Styles gallery."""
    el = style.element
    # qFormat lives after w:name/w:basedOn/w:next/w:link/w:uiPriority and
    # before run/paragraph property blocks. Appending works in practice
    # because Word tolerates trailing flag elements before rPr/pPr; to be
    # safe, insert immediately after w:name if present.
    qfmt = OxmlElement("w:qFormat")
    name = el.find(qn("w:name"))
    if name is not None:
        name.addnext(qfmt)
    else:
        el.append(qfmt)


def _ensure_styles(doc) -> None:
    """Create the minimal custom style set used for round-trippable output."""
    styles = doc.styles
    existing = {s.name for s in styles}

    # Inline code character style.
    if INLINE_CODE_STYLE_NAME not in existing:
        s = styles.add_style(INLINE_CODE_STYLE_NAME, WD_STYLE_TYPE.CHARACTER)
        s.font.name = CODE_FONT
        s.font.size = Pt(10)
        r_pr = s.element.get_or_add_rPr()
        r_pr.append(_shading(CODE_SHADING_FILL))
        # Force fixed-width font for all script ranges.
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), CODE_FONT)
        rfonts.set(qn("w:hAnsi"), CODE_FONT)
        rfonts.set(qn("w:cs"), CODE_FONT)
        r_pr.append(rfonts)
        _mark_quick_format(s)

    # Code block paragraph style.
    if CODE_BLOCK_STYLE_NAME not in existing:
        s = styles.add_style(CODE_BLOCK_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = CODE_FONT
        s.font.size = Pt(10)
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(0)
        p_pr = s.element.get_or_add_pPr()
        p_pr.append(_shading(CODE_SHADING_FILL))
        _mark_quick_format(s)


# ---------------------------------------------------------------------------
# Markdown -> DOCX conversion
# ---------------------------------------------------------------------------

class InlineFormat:
    """Mutable bag of inline formatting flags as we walk inline tokens."""

    __slots__ = ("bold", "italic", "strike", "code")

    def __init__(self) -> None:
        self.bold = False
        self.italic = False
        self.strike = False
        self.code = False

    def copy(self) -> "InlineFormat":
        c = InlineFormat()
        c.bold = self.bold
        c.italic = self.italic
        c.strike = self.strike
        c.code = self.code
        return c


def _apply_run_formatting(run, fmt: InlineFormat,
                          inline_code_style) -> None:
    if fmt.code:
        run.style = inline_code_style
    if fmt.bold:
        run.bold = True
    if fmt.italic:
        run.italic = True
    if fmt.strike:
        run.font.strike = True


def _wrap_runs_in_hyperlink(paragraph, runs, url: str) -> None:
    """Move trailing w:r elements into a new w:hyperlink with relationship."""
    if not runs:
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    first_r = runs[0]._element
    first_r.addprevious(hyperlink)
    for r in runs:
        paragraph._p.remove(r._element)
        hyperlink.append(r._element)

    # Apply Hyperlink character style to each run.
    for r in runs:
        r_pr = r._element.get_or_add_rPr()
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "Hyperlink")
        r_pr.insert(0, rstyle)


def _render_inline(doc, paragraph, tokens: List[Token],
                   fmt: InlineFormat, inline_code_style) -> None:
    """Render an inline token stream into the given paragraph."""
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        t = tok.type

        if t == "text":
            run = paragraph.add_run(tok.content)
            _apply_run_formatting(run, fmt, inline_code_style)

        elif t == "code_inline":
            sub = fmt.copy()
            sub.code = True
            run = paragraph.add_run(tok.content)
            _apply_run_formatting(run, sub, inline_code_style)

        elif t == "softbreak":
            run = paragraph.add_run(" ")
            _apply_run_formatting(run, fmt, inline_code_style)

        elif t == "hardbreak":
            run = paragraph.add_run()
            run.add_break()

        elif t == "strong_open":
            fmt.bold = True
        elif t == "strong_close":
            fmt.bold = False

        elif t == "em_open":
            fmt.italic = True
        elif t == "em_close":
            fmt.italic = False

        elif t == "s_open":
            fmt.strike = True
        elif t == "s_close":
            fmt.strike = False

        elif t == "link_open":
            href = tok.attrGet("href") or ""
            # Find matching link_close, recurse on inner tokens, then wrap.
            depth = 1
            j = i + 1
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "link_open":
                    depth += 1
                elif tokens[j].type == "link_close":
                    depth -= 1
                if depth == 0:
                    break
                j += 1
            inner_tokens = tokens[i + 1:j]

            # Snapshot existing runs; render inner; identify new runs; wrap.
            existing_run_elements = set(id(r._element) for r in paragraph.runs)
            _render_inline(doc, paragraph, inner_tokens, fmt, inline_code_style)
            new_runs = [r for r in paragraph.runs
                        if id(r._element) not in existing_run_elements]
            if href:
                _wrap_runs_in_hyperlink(paragraph, new_runs, href)
            i = j  # skip past link_close

        elif t == "link_close":
            # Already consumed by link_open branch.
            pass

        elif t == "image":
            # v1: images are out of scope; emit alt text as plain run.
            alt = tok.content or tok.attrGet("alt") or ""
            if alt:
                run = paragraph.add_run(alt)
                _apply_run_formatting(run, fmt, inline_code_style)

        elif t == "html_inline":
            # v1: raw HTML is out of scope; pass through as literal text so
            # nothing is silently dropped.
            run = paragraph.add_run(tok.content)
            _apply_run_formatting(run, fmt, inline_code_style)

        else:
            # Unknown inline token. Fail fast per spec.
            fail(f"Unsupported inline markdown token: {t}")

        i += 1


def _cell_alignment(style_attr: str) -> Optional[int]:
    if not style_attr:
        return None
    if "text-align:right" in style_attr:
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    if "text-align:center" in style_attr:
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    if "text-align:left" in style_attr:
        return WD_PARAGRAPH_ALIGNMENT.LEFT
    return None


class BlockRenderer:
    """Walks a markdown-it block token stream and emits DOCX content."""

    def __init__(self, doc) -> None:
        self.doc = doc
        self.inline_code_style = doc.styles[INLINE_CODE_STYLE_NAME]
        # Each entry: {"ordered": bool, "level": int}
        self.list_stack: List[Dict[str, object]] = []
        # Set when a list_item_open is processed; consumed by the next
        # paragraph created. Cleared after first use so subsequent paragraphs
        # inside the same item become continuation paragraphs.
        self._pending_list_item = False
        # Depth of enclosing blockquotes; non-zero means paragraphs not
        # otherwise styled (e.g. by list state) get the Quote style.
        self._blockquote_depth = 0

    # -- entry point --

    def render(self, tokens: List[Token]) -> None:
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            i = self._render_block(tokens, i, tok)

    # -- helpers --

    def _find_matching_close(self, tokens: List[Token], start: int,
                             open_type: str, close_type: str) -> int:
        depth = 1
        j = start + 1
        while j < len(tokens):
            if tokens[j].type == open_type:
                depth += 1
            elif tokens[j].type == close_type:
                depth -= 1
                if depth == 0:
                    return j
            j += 1
        fail(f"Malformed markdown: missing {close_type}")
        return -1  # unreachable

    def _inline_tokens_of(self, tokens: List[Token], idx: int) -> List[Token]:
        """If tokens[idx] is an inline token, return its children list."""
        if idx < len(tokens) and tokens[idx].type == "inline":
            return tokens[idx].children or []
        return []

    # -- block dispatch --

    def _render_block(self, tokens: List[Token], i: int, tok: Token) -> int:
        t = tok.type

        if t == "heading_open":
            level = int(tok.tag[1])  # 'h1' -> 1
            close_idx = self._find_matching_close(
                tokens, i, "heading_open", "heading_close"
            )
            inline_idx = i + 1
            p = self.doc.add_paragraph(style=f"Heading {level}")
            fmt = InlineFormat()
            _render_inline(
                self.doc, p, self._inline_tokens_of(tokens, inline_idx),
                fmt, self.inline_code_style,
            )
            return close_idx + 1

        if t == "paragraph_open":
            close_idx = self._find_matching_close(
                tokens, i, "paragraph_open", "paragraph_close"
            )
            inline_idx = i + 1
            p = self._new_paragraph()
            fmt = InlineFormat()
            _render_inline(
                self.doc, p, self._inline_tokens_of(tokens, inline_idx),
                fmt, self.inline_code_style,
            )
            return close_idx + 1

        if t == "bullet_list_open":
            close_idx = self._find_matching_close(
                tokens, i, "bullet_list_open", "bullet_list_close"
            )
            self.list_stack.append({
                "ordered": False,
                "level": len(self.list_stack),
            })
            self._render_list_body(tokens, i + 1, close_idx)
            self.list_stack.pop()
            return close_idx + 1

        if t == "ordered_list_open":
            close_idx = self._find_matching_close(
                tokens, i, "ordered_list_open", "ordered_list_close"
            )
            self.list_stack.append({
                "ordered": True,
                "level": len(self.list_stack),
            })
            self._render_list_body(tokens, i + 1, close_idx)
            self.list_stack.pop()
            return close_idx + 1

        if t == "blockquote_open":
            close_idx = self._find_matching_close(
                tokens, i, "blockquote_open", "blockquote_close"
            )
            self._render_blockquote(tokens, i + 1, close_idx)
            return close_idx + 1

        if t == "fence" or t == "code_block":
            self._render_code_block(tok.content)
            return i + 1

        if t == "hr":
            self._render_hr()
            return i + 1

        if t == "table_open":
            close_idx = self._find_matching_close(
                tokens, i, "table_open", "table_close"
            )
            self._render_table(tokens, i + 1, close_idx)
            return close_idx + 1

        if t == "html_block":
            # v1: raw HTML blocks are unsupported. Spec marks them as such.
            # Emit content as a plain paragraph so nothing is silently lost,
            # but flag this loudly.
            sys.stderr.write(
                "WARNING: raw HTML block encountered; emitting as plain text.\n"
            )
            p = self.doc.add_paragraph(tok.content.rstrip("\n"))
            return i + 1

        if t in ("paragraph_close", "heading_close", "bullet_list_close",
                 "ordered_list_close", "list_item_close", "blockquote_close",
                 "table_close", "thead_close", "tbody_close", "tr_close",
                 "th_close", "td_close"):
            return i + 1

        # Unknown block token. Fail fast.
        fail(f"Unsupported block markdown token: {t}")
        return i + 1  # unreachable

    # -- block helpers --

    def _new_paragraph(self):
        """Create a paragraph, applying the correct style for the current
        block context (list item, blockquote, or plain)."""
        # Highest priority: a list item just opened; consume the slot.
        if self._pending_list_item and self.list_stack:
            self._pending_list_item = False
            top = self.list_stack[-1]
            style_name = "List Number" if top["ordered"] else "List Bullet"
            p = self.doc.add_paragraph(style=style_name)
            level = int(top["level"])
            if level > 0:
                # Built-in template does not guarantee "List Bullet 2" etc.,
                # so we keep the base style and apply manual indentation.
                p.paragraph_format.left_indent = Cm(0.75 + 0.75 * level)
            return p

        # Continuation paragraph inside an already-bulleted list item.
        if self.list_stack:
            p = self.doc.add_paragraph()
            level = len(self.list_stack)
            p.paragraph_format.left_indent = Cm(0.75 + 0.75 * max(level - 1, 0) + 0.6)
            return p

        # Blockquote with no enclosing list.
        if self._blockquote_depth > 0:
            try:
                return self.doc.add_paragraph(style=QUOTE_STYLE_NAME)
            except KeyError:
                return self.doc.add_paragraph()

        return self.doc.add_paragraph()

    def _render_list_body(self, tokens: List[Token],
                          start: int, end_exclusive: int) -> None:
        i = start
        while i < end_exclusive:
            tok = tokens[i]
            if tok.type == "list_item_open":
                close_idx = self._find_matching_close(
                    tokens, i, "list_item_open", "list_item_close"
                )
                self._render_list_item(tokens, i + 1, close_idx)
                i = close_idx + 1
            else:
                i += 1

    def _render_list_item(self, tokens: List[Token],
                          start: int, end_exclusive: int) -> None:
        """Render one list item. The first block-level paragraph emitted
        consumes the list bullet; subsequent paragraphs in the same item
        become continuation paragraphs (indented, no bullet)."""
        self._pending_list_item = True
        try:
            i = start
            while i < end_exclusive:
                i = self._render_block(tokens, i, tokens[i])
        finally:
            # If the item had no paragraph content at all (rare), clear flag
            # so it does not leak into a sibling.
            self._pending_list_item = False

    def _render_blockquote(self, tokens: List[Token],
                           start: int, end_exclusive: int) -> None:
        self._blockquote_depth += 1
        try:
            i = start
            while i < end_exclusive:
                i = self._render_block(tokens, i, tokens[i])
        finally:
            self._blockquote_depth -= 1

    def _render_code_block(self, content: str) -> None:
        # Strip a single trailing newline; preserve internal newlines as
        # separate paragraphs all sharing the code block style.
        text = content.rstrip("\n")
        if text == "":
            self.doc.add_paragraph(style=CODE_BLOCK_STYLE_NAME)
            return
        for line in text.split("\n"):
            p = self.doc.add_paragraph(style=CODE_BLOCK_STYLE_NAME)
            if line:
                p.add_run(line)

    def _render_hr(self) -> None:
        p = self.doc.add_paragraph()
        _add_paragraph_bottom_border(p)

    def _render_table(self, tokens: List[Token],
                      start: int, end_exclusive: int) -> None:
        # First pass: collect rows of cells with alignment metadata.
        rows: List[List[Tuple[List[Token], Optional[int], bool]]] = []
        # cell = (inline_tokens, alignment, is_header)

        i = start
        while i < end_exclusive:
            tok = tokens[i]
            if tok.type == "tr_open":
                tr_close = self._find_matching_close(
                    tokens, i, "tr_open", "tr_close"
                )
                row: List[Tuple[List[Token], Optional[int], bool]] = []
                j = i + 1
                while j < tr_close:
                    ctok = tokens[j]
                    if ctok.type in ("th_open", "td_open"):
                        is_header = (ctok.type == "th_open")
                        cell_close = self._find_matching_close(
                            tokens, j,
                            "th_open" if is_header else "td_open",
                            "th_close" if is_header else "td_close",
                        )
                        style_attr = ctok.attrGet("style") or ""
                        alignment = _cell_alignment(style_attr)
                        inline_tokens: List[Token] = []
                        k = j + 1
                        while k < cell_close:
                            if tokens[k].type == "inline":
                                inline_tokens = tokens[k].children or []
                                break
                            k += 1
                        row.append((inline_tokens, alignment, is_header))
                        j = cell_close + 1
                    else:
                        j += 1
                rows.append(row)
                i = tr_close + 1
            else:
                i += 1

        if not rows:
            return

        ncols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        try:
            table.style = TABLE_STYLE_NAME
        except KeyError:
            pass

        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell = table.rows[ri].cells[ci]
                # Clear the default empty paragraph python-docx inserts.
                # We will write our content into the first paragraph directly.
                target_p = cell.paragraphs[0]
                target_p.text = ""
                if ci < len(row):
                    inline_tokens, alignment, is_header = row[ci]
                    fmt = InlineFormat()
                    if is_header:
                        fmt.bold = True
                    _render_inline(
                        self.doc, target_p, inline_tokens,
                        fmt, self.inline_code_style,
                    )
                    if alignment is not None:
                        target_p.alignment = alignment


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert one markdown file to a docx working copy. Fail fast on error."""
    try:
        with open(md_path, "r", encoding="utf-8") as fp:
            md_text = fp.read()
    except OSError as exc:
        fail(f"Failed to read {md_path}: {exc}")
        return  # unreachable
    except UnicodeDecodeError as exc:
        fail(f"{md_path} is not valid UTF-8: {exc}")
        return  # unreachable

    md = MarkdownIt("commonmark", {"html": False, "linkify": False, "breaks": False})
    md.enable(["table", "strikethrough"])

    try:
        tokens = md.parse(md_text)
    except Exception as exc:  # markdown-it rarely raises, but be defensive.
        fail(f"Failed to parse {md_path}: {exc}")
        return  # unreachable

    doc = Document()
    _ensure_styles(doc)

    renderer = BlockRenderer(doc)
    renderer.render(tokens)

    try:
        doc.save(str(docx_path))
    except OSError as exc:
        fail(f"Failed to write {docx_path}: {exc}")


# ---------------------------------------------------------------------------
# --togit: DOCX -> Markdown
# ---------------------------------------------------------------------------
#
# Design notes:
#
# * Style detection is anchored to the styles --fromgit produces. The user is
#   expected to edit content within that style vocabulary; styles outside the
#   supported set are mapped to plain paragraphs (per spec: "do not preserve
#   arbitrary Word formatting outside the supported markdown-compatible style
#   set").
# * Block order: walk document.element.body children directly so paragraphs
#   and tables appear in their actual document order (python-docx exposes
#   them as separate lists which loses interleaving).
# * Tracked changes: detected via <w:ins>/<w:del> presence in the document
#   XML. python-docx does not cleanly expose an accepted-view, so per spec
#   ("if tracked changes cannot be resolved safely and deterministically,
#   fail fast"), refuse to convert until the user accepts/rejects them.
# * Images: rendered at IMAGE_DPI of their Word display size, never upscaled.
#   Hashed by their rendered bytes. Filenames are stable across runs when
#   content is unchanged; changed content gets a new filename.

# --- OOXML namespace shortcuts used in this section ---
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}

MC_FALLBACK_TAG = f"{{{NS['mc']}}}Fallback"
MC_CHOICE_TAG = f"{{{NS['mc']}}}Choice"


# ---------------------------------------------------------------------------
# Manifest
# ---------------------------------------------------------------------------

def _manifest_path(git_dir: Path) -> Path:
    return git_dir / ASSETS_DIR_NAME / MANIFEST_FILENAME


def load_manifest(git_dir: Path) -> Dict:
    path = _manifest_path(git_dir)
    if not path.exists():
        return {"version": MANIFEST_VERSION, "documents": {}}
    try:
        with open(path, "r", encoding="utf-8") as fp:
            data = json.load(fp)
    except (OSError, json.JSONDecodeError) as exc:
        fail(f"Manifest is corrupt or unreadable ({path}): {exc}")
        return {}  # unreachable
    if not isinstance(data, dict) or "documents" not in data:
        fail(f"Manifest has unexpected shape: {path}")
    data.setdefault("version", MANIFEST_VERSION)
    data.setdefault("documents", {})
    return data


def save_manifest(git_dir: Path, manifest: Dict) -> None:
    path = _manifest_path(git_dir)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    try:
        with open(tmp, "w", encoding="utf-8") as fp:
            json.dump(manifest, fp, indent=2, sort_keys=True)
            fp.write("\n")
        os.replace(tmp, path)
    except OSError as exc:
        try:
            if tmp.exists():
                tmp.unlink()
        except OSError:
            pass
        fail(f"Failed to update manifest {path}: {exc}")


# ---------------------------------------------------------------------------
# Sync state (per-WorkDir change detection)
# ---------------------------------------------------------------------------
#
# Records mtime + size of each .docx at the time of its last successful
# conversion. A subsequent --togit run skips files whose mtime AND size
# match the recorded values. mtime+size combined is cheap, deterministic,
# and good enough in practice; a content hash would catch the (rare) case
# of an edit that preserves both, at the cost of hashing every file each
# run.
#
# Deleting this file forces a full resync on the next run.

def _state_path(work_dir: Path) -> Path:
    return work_dir / STATE_FILENAME


def load_state(work_dir: Path) -> Dict:
    path = _state_path(work_dir)
    if not path.exists():
        return {"version": STATE_VERSION, "files": {}}
    try:
        with open(path, "r", encoding="utf-8") as fp:
            data = json.load(fp)
    except (OSError, json.JSONDecodeError) as exc:
        fail(f"Sync state file is corrupt or unreadable ({path}): {exc}")
        return {}  # unreachable
    if not isinstance(data, dict) or "files" not in data:
        fail(f"Sync state file has unexpected shape: {path}")
    data.setdefault("version", STATE_VERSION)
    data.setdefault("files", {})
    return data


def save_state(work_dir: Path, state: Dict) -> None:
    path = _state_path(work_dir)
    tmp = path.with_suffix(path.suffix + ".tmp")
    try:
        with open(tmp, "w", encoding="utf-8") as fp:
            json.dump(state, fp, indent=2, sort_keys=True)
            fp.write("\n")
        os.replace(tmp, path)
    except OSError as exc:
        try:
            if tmp.exists():
                tmp.unlink()
        except OSError:
            pass
        fail(f"Failed to update sync state {path}: {exc}")


def _docx_unchanged_since_last_sync(docx_path: Path, state: Dict) -> bool:
    entry = state["files"].get(docx_path.name)
    if not entry:
        return False
    try:
        stat = docx_path.stat()
    except OSError:
        return False
    return (
        stat.st_mtime == entry.get("mtime")
        and stat.st_size == entry.get("size")
    )


def _record_sync(docx_path: Path, state: Dict) -> None:
    stat = docx_path.stat()
    state["files"][docx_path.name] = {
        "mtime": stat.st_mtime,
        "size": stat.st_size,
        "last_synced": datetime.now(timezone.utc).strftime(
            "%Y-%m-%dT%H:%M:%SZ"
        ),
    }


# ---------------------------------------------------------------------------
# Page number tracking
# ---------------------------------------------------------------------------
#
# OOXML does not store explicit page numbers; layout is dynamic. Word does
# insert <w:lastRenderedPageBreak/> markers in runs after each layout pass,
# and explicit page breaks appear as <w:br w:type="page"/>. Counting both
# in document order up to a target element gives a page number that is
# accurate as of the last time Word laid out the document.
#
# Build a single map once at the start of a conversion and look up element
# pages from it. Cheap, deterministic, and avoids quadratic walks.

class PageTracker:
    def __init__(self, doc) -> None:
        # Map id(element) -> page number for every body descendant.
        self._page_by_id: Dict[int, int] = {}
        page = 1
        for el in _iter_skipping_fallback(doc.element.body):
            self._page_by_id[id(el)] = page
            tag = el.tag
            if tag == qn("w:lastRenderedPageBreak"):
                page += 1
            elif tag == qn("w:br") and el.get(qn("w:type")) == "page":
                page += 1

    def page_for(self, element) -> int:
        if element is None:
            return 1
        # Walk up until we find an element we have a page for.
        cur = element
        while cur is not None:
            pg = self._page_by_id.get(id(cur))
            if pg is not None:
                return pg
            cur = cur.getparent()
        return 1


def _iter_skipping_fallback(root):
    """Iterate root and all descendants, but skip subtrees inside
    <mc:Fallback>. The fallback is the legacy VML version of content that
    is also present in DrawingML inside <mc:Choice>; iterating both would
    double-count shapes and confuse drawing detection.
    """
    if root.tag == MC_FALLBACK_TAG:
        return
    yield root
    for child in root:
        yield from _iter_skipping_fallback(child)


def _iter_drawings_in(element):
    """Yield <w:drawing> descendants of element in document order,
    excluding any inside <mc:Fallback>."""
    for el in _iter_skipping_fallback(element):
        if el.tag == qn("w:drawing"):
            yield el


def _fail_with_page(message: str, page: Optional[int]) -> None:
    suffix = f" (page {page})" if page else ""
    fail(message + suffix)


# ---------------------------------------------------------------------------
# Floating shape parsing
# ---------------------------------------------------------------------------
#
# v1 supports a constrained subset of DrawingML shapes that commonly serve
# as annotation overlays on screenshots:
#
#   - rect          : rectangle (optional fill, optional outline)
#   - line          : straight line
#   - straightConnector1..5 : straight connector lines (lines with optional
#     arrowheads at either end)
#   - ellipse       : ellipse / oval
#
# Shapes carrying text are supported when their preset geometry is rect or
# ellipse (text-only callouts not supported).
#
# Anything outside this set raises fail-fast with the shape type and page.

SUPPORTED_SHAPE_GEOMETRY = {
    "rect", "ellipse",
    "line",
    "straightConnector1", "straightConnector2", "straightConnector3",
    "straightConnector4", "straightConnector5",
}
LINE_GEOMETRIES = {
    "line",
    "straightConnector1", "straightConnector2", "straightConnector3",
    "straightConnector4", "straightConnector5",
}


class FloatingShape:
    """A parsed floating shape, in EMU coordinates relative to the page-
    level container the anchor uses (we treat positionH/V offsets as the
    shape's position; the caller decides whether that position is contained
    within the inline image's extent)."""

    __slots__ = (
        "kind", "x_emu", "y_emu", "width_emu", "height_emu",
        "fill_rgb", "stroke_rgb", "stroke_width_emu",
        "flip_h", "flip_v", "rotation_60k",
        "head_end_present", "tail_end_present",
        "text_runs",
        "text_anchor", "text_ins_l", "text_ins_t", "text_ins_r", "text_ins_b",
    )

    def __init__(self):
        self.kind: str = ""
        self.x_emu: int = 0
        self.y_emu: int = 0
        self.width_emu: int = 0
        self.height_emu: int = 0
        self.fill_rgb: Optional[Tuple[int, int, int]] = None
        self.stroke_rgb: Optional[Tuple[int, int, int]] = None
        self.stroke_width_emu: int = 0
        self.flip_h: bool = False
        self.flip_v: bool = False
        self.rotation_60k: int = 0  # 60000ths of a degree, OOXML convention
        self.head_end_present: bool = False
        self.tail_end_present: bool = False
        # text_runs: list of (text, {bold, italic, size_pt, rgb})
        self.text_runs: List[Tuple[str, Dict]] = []
        # bodyPr properties. Defaults match Word's defaults when unspecified.
        self.text_anchor: str = "t"  # 't' (top), 'ctr' (centre), 'b' (bottom)
        self.text_ins_l: int = 91440   # 0.1 inch
        self.text_ins_t: int = 45720   # 0.05 inch
        self.text_ins_r: int = 91440
        self.text_ins_b: int = 45720


def _parse_srgb(el) -> Optional[Tuple[int, int, int]]:
    """Extract an sRGB triple from <a:srgbClr val="HHHHHH"/> inside el."""
    if el is None:
        return None
    clr = el.find(f".//{{{NS['a']}}}srgbClr")
    if clr is None:
        return None
    val = clr.get("val", "")
    if len(val) != 6:
        return None
    try:
        return (int(val[0:2], 16), int(val[2:4], 16), int(val[4:6], 16))
    except ValueError:
        return None


def _parse_shape_text(wsp_el) -> List[Tuple[str, Dict]]:
    """Extract text runs from a <wps:txbx> child of a wps:wsp shape."""
    txbx = wsp_el.find(f"{{{NS['wps']}}}txbx")
    if txbx is None:
        return []
    content = txbx.find(f"{{{NS['w']}}}txbxContent")
    if content is None:
        return []
    out: List[Tuple[str, Dict]] = []
    for r in content.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        fmt: Dict = {"bold": False, "italic": False, "size_pt": None,
                     "rgb": None}
        if rPr is not None:
            if rPr.find(qn("w:b")) is not None:
                fmt["bold"] = True
            if rPr.find(qn("w:i")) is not None:
                fmt["italic"] = True
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                try:
                    # w:sz val is in half-points.
                    fmt["size_pt"] = int(sz.get(qn("w:val"), "0")) / 2.0
                except ValueError:
                    pass
            color = rPr.find(qn("w:color"))
            if color is not None:
                cval = color.get(qn("w:val"), "")
                if len(cval) == 6:
                    try:
                        fmt["rgb"] = (
                            int(cval[0:2], 16),
                            int(cval[2:4], 16),
                            int(cval[4:6], 16),
                        )
                    except ValueError:
                        pass
        parts: List[str] = []
        for child in r:
            if child.tag == qn("w:t") and child.text:
                parts.append(child.text)
            elif child.tag == qn("w:tab"):
                parts.append("\t")
            elif child.tag == qn("w:br"):
                parts.append("\n")
        text = "".join(parts)
        if text:
            out.append((text, fmt))
    return out


def _parse_anchor_drawing(drawing_el, page: int, docx_path: Path
                          ) -> Optional[FloatingShape]:
    """Parse a <w:drawing> with <wp:anchor>. Returns None if the anchor
    holds a picture (handled by image extraction); otherwise returns the
    parsed FloatingShape. Fails fast for unsupported shape kinds."""
    anchor = drawing_el.find(f"{{{NS['wp']}}}anchor")
    if anchor is None:
        return None  # not an anchored drawing

    # If the anchor's graphic data is a picture, this is a floating image,
    # not an annotation shape. We surface it as the existing fail-fast in
    # the image-extraction code path; here we just return None so the
    # caller can detect "contains picture" via separate logic.
    graphic_data = anchor.find(
        f".//{{{NS['a']}}}graphicData"
    )
    if graphic_data is None:
        _fail_with_page(
            f"{docx_path.name} contains an anchored drawing with no graphic "
            f"data.",
            page,
        )
        return None  # unreachable

    uri = graphic_data.get("uri", "")
    if uri == "http://schemas.openxmlformats.org/drawingml/2006/picture":
        # Floating picture (image). Distinct case from shape annotations:
        # the existing image-extraction fail-fast covers it.
        return None

    if uri != "http://schemas.microsoft.com/office/word/2010/wordprocessingShape":
        _fail_with_page(
            f"{docx_path.name} contains an unsupported anchored drawing "
            f"type ({uri!r}). Only word-processing shapes are supported "
            f"as image annotations.",
            page,
        )
        return None  # unreachable

    wsp = graphic_data.find(f"{{{NS['wps']}}}wsp")
    if wsp is None:
        _fail_with_page(
            f"{docx_path.name} contains a shape with no wps:wsp body.",
            page,
        )
        return None  # unreachable

    shape = FloatingShape()

    # Position: <wp:positionH><wp:posOffset>, <wp:positionV><wp:posOffset>.
    pos_h = anchor.find(f"{{{NS['wp']}}}positionH")
    pos_v = anchor.find(f"{{{NS['wp']}}}positionV")
    if pos_h is None or pos_v is None:
        _fail_with_page(
            f"{docx_path.name} contains an anchored shape without explicit "
            f"position offsets; only absolute-offset positioning is "
            f"supported.",
            page,
        )
        return None  # unreachable
    off_h = pos_h.find(f"{{{NS['wp']}}}posOffset")
    off_v = pos_v.find(f"{{{NS['wp']}}}posOffset")
    if off_h is None or off_v is None or off_h.text is None or off_v.text is None:
        _fail_with_page(
            f"{docx_path.name} contains an anchored shape positioned by "
            f"alignment rather than absolute offset, which cannot be "
            f"reliably positioned over an inline image.",
            page,
        )
        return None  # unreachable
    try:
        shape.x_emu = int(off_h.text)
        shape.y_emu = int(off_v.text)
    except ValueError:
        _fail_with_page(
            f"{docx_path.name} contains an anchored shape with non-numeric "
            f"position offset.",
            page,
        )
        return None  # unreachable

    extent = anchor.find(f"{{{NS['wp']}}}extent")
    if extent is None:
        _fail_with_page(
            f"{docx_path.name} contains an anchored shape with no extent.",
            page,
        )
        return None  # unreachable
    try:
        shape.width_emu = int(extent.get("cx", "0"))
        shape.height_emu = int(extent.get("cy", "0"))
    except ValueError:
        _fail_with_page(
            f"{docx_path.name} contains an anchored shape with non-numeric "
            f"extent.",
            page,
        )
        return None  # unreachable

    sp_pr = wsp.find(f"{{{NS['wps']}}}spPr")
    if sp_pr is None:
        _fail_with_page(
            f"{docx_path.name} contains a shape with no spPr properties.",
            page,
        )
        return None  # unreachable

    # Transform: flip/rotation from <a:xfrm>.
    xfrm = sp_pr.find(f"{{{NS['a']}}}xfrm")
    if xfrm is not None:
        shape.flip_h = xfrm.get("flipH", "0") == "1"
        shape.flip_v = xfrm.get("flipV", "0") == "1"
        try:
            shape.rotation_60k = int(xfrm.get("rot", "0"))
        except ValueError:
            shape.rotation_60k = 0

    if shape.rotation_60k % (360 * 60000) != 0:
        _fail_with_page(
            f"{docx_path.name} contains a rotated shape (rotation is not "
            f"supported for image annotations in v1).",
            page,
        )
        return None  # unreachable

    # Geometry: <a:prstGeom prst="...">.
    prst_geom = sp_pr.find(f"{{{NS['a']}}}prstGeom")
    if prst_geom is None:
        _fail_with_page(
            f"{docx_path.name} contains a shape without a preset geometry; "
            f"only preset shapes (rect, ellipse, line, connectors) are "
            f"supported.",
            page,
        )
        return None  # unreachable
    prst = prst_geom.get("prst", "")
    if prst not in SUPPORTED_SHAPE_GEOMETRY:
        _fail_with_page(
            f"{docx_path.name} contains a shape with unsupported geometry "
            f"{prst!r}. Supported: rectangle, ellipse, straight line, "
            f"straight connectors.",
            page,
        )
        return None  # unreachable
    shape.kind = "line" if prst in LINE_GEOMETRIES else prst

    # Fill: <a:solidFill> for solid; <a:noFill/> means no fill.
    if sp_pr.find(f"{{{NS['a']}}}noFill") is None:
        solid = sp_pr.find(f"{{{NS['a']}}}solidFill")
        if solid is not None:
            shape.fill_rgb = _parse_srgb(solid)
            if shape.fill_rgb is None:
                # Gradient/pattern/scheme color: degrade silently to no fill.
                # Documented v1 limitation.
                pass

    # Stroke: <a:ln w="EMU">[<a:solidFill>][<a:headEnd/>][<a:tailEnd/>]</a:ln>.
    ln = sp_pr.find(f"{{{NS['a']}}}ln")
    if ln is not None:
        try:
            shape.stroke_width_emu = int(ln.get("w", "0"))
        except ValueError:
            shape.stroke_width_emu = 0
        if ln.find(f"{{{NS['a']}}}noFill") is None:
            shape.stroke_rgb = _parse_srgb(ln)
        head_end = ln.find(f"{{{NS['a']}}}headEnd")
        tail_end = ln.find(f"{{{NS['a']}}}tailEnd")
        if head_end is not None and head_end.get("type", "none") != "none":
            shape.head_end_present = True
        if tail_end is not None and tail_end.get("type", "none") != "none":
            shape.tail_end_present = True
    # Lines without an explicit ln default to a thin black stroke in Word.
    if shape.kind == "line" and shape.stroke_rgb is None and ln is None:
        shape.stroke_rgb = (0, 0, 0)
        shape.stroke_width_emu = 9525  # ~0.75pt

    # Text content for text boxes.
    if shape.kind in ("rect", "ellipse"):
        shape.text_runs = _parse_shape_text(wsp)
        # bodyPr governs how text sits inside the shape: vertical anchor
        # and per-edge insets. Missing values fall back to Word's defaults.
        body_pr = wsp.find(f"{{{NS['wps']}}}bodyPr")
        if body_pr is not None:
            anchor = body_pr.get("anchor", "t")
            if anchor in ("t", "ctr", "b"):
                shape.text_anchor = anchor
            for attr, field in (("lIns", "text_ins_l"),
                                ("tIns", "text_ins_t"),
                                ("rIns", "text_ins_r"),
                                ("bIns", "text_ins_b")):
                val = body_pr.get(attr)
                if val is not None:
                    try:
                        setattr(shape, field, int(val))
                    except ValueError:
                        pass
    else:
        text_runs = _parse_shape_text(wsp)
        if text_runs:
            _fail_with_page(
                f"{docx_path.name} contains a non-rectangular text-bearing "
                f"shape, which v1 does not support.",
                page,
            )
            return None  # unreachable

    return shape


def _inline_image_extent(drawing_el, page: int,
                         docx_path: Path) -> Tuple[int, int]:
    """Return (cx, cy) EMU for an inline image. Caller has verified inline."""
    inline = drawing_el.find(f"{{{NS['wp']}}}inline")
    extent = inline.find(f"{{{NS['wp']}}}extent")
    if extent is None:
        _fail_with_page(
            f"{docx_path.name} has an inline image with no display extent.",
            page,
        )
        return (0, 0)  # unreachable
    try:
        return (int(extent.get("cx", "0")), int(extent.get("cy", "0")))
    except ValueError:
        _fail_with_page(
            f"{docx_path.name} has an inline image with non-numeric extent.",
            page,
        )
        return (0, 0)  # unreachable


def _shape_within_extent(shape: FloatingShape,
                         image_extent: Tuple[int, int]) -> bool:
    """Whether the shape's full bounding box lies inside an image whose
    origin is treated as (0, 0) and extent is image_extent (in EMU).

    A small tolerance handles rounding noise from Word."""
    cx, cy = image_extent
    tol = 9525  # ~0.75 point (one stroke width) of slack
    return (
        shape.x_emu >= -tol
        and shape.y_emu >= -tol
        and (shape.x_emu + shape.width_emu) <= cx + tol
        and (shape.y_emu + shape.height_emu) <= cy + tol
    )


# ---------------------------------------------------------------------------
# Shape rendering with Pillow
# ---------------------------------------------------------------------------

def _emu_to_px_f(emu: float, dpi: int = IMAGE_DPI) -> float:
    return emu * dpi / 914400.0


def _pick_font(size_px: int):
    """Pick a TrueType font path. Tries common Windows fonts first
    (since the tool targets Windows), then Linux/macOS fallbacks, then
    Pillow's default bitmap font as last resort."""
    from PIL import ImageFont
    candidates = [
        "C:\\Windows\\Fonts\\arial.ttf",
        "C:\\Windows\\Fonts\\calibri.ttf",
        "C:\\Windows\\Fonts\\segoeui.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=max(8, int(size_px)))
        except (OSError, IOError):
            continue
    return ImageFont.load_default()


def _composite_shapes(base_img, shapes: List[FloatingShape],
                      image_extent_emu: Tuple[int, int]) -> "Image.Image":
    """Render shapes on top of base_img (which is the inline image already
    scaled to its Word display size at IMAGE_DPI). All shape coordinates
    are interpreted as offsets in EMU from the image origin."""
    from PIL import Image as PILImage, ImageDraw

    # Scale factor from EMU to pixels in the rendered image. The rendered
    # image's pixel width may be slightly smaller than display-size at
    # IMAGE_DPI because we never upscale beyond the original; compute the
    # actual scale.
    img_cx_emu, img_cy_emu = image_extent_emu
    if img_cx_emu <= 0 or img_cy_emu <= 0:
        return base_img
    scale_x = base_img.width / img_cx_emu
    scale_y = base_img.height / img_cy_emu

    def px_x(emu: float) -> float:
        return emu * scale_x

    def px_y(emu: float) -> float:
        return emu * scale_y

    # Composite onto an RGBA canvas to handle transparency cleanly.
    if base_img.mode != "RGBA":
        canvas = base_img.convert("RGBA")
    else:
        canvas = base_img.copy()
    draw = ImageDraw.Draw(canvas, "RGBA")

    for shape in shapes:
        if shape.kind == "rect":
            _draw_rect_shape(draw, canvas, shape, px_x, px_y)
        elif shape.kind == "ellipse":
            _draw_ellipse_shape(draw, canvas, shape, px_x, px_y)
        elif shape.kind == "line":
            _draw_line_shape(draw, shape, px_x, px_y)
        # Unsupported kinds are rejected upstream.

    return canvas


def _draw_rect_shape(draw, canvas, shape: FloatingShape,
                     px_x, px_y) -> None:
    x0 = px_x(shape.x_emu)
    y0 = px_y(shape.y_emu)
    x1 = px_x(shape.x_emu + shape.width_emu)
    y1 = px_y(shape.y_emu + shape.height_emu)
    if x1 < x0:
        x0, x1 = x1, x0
    if y1 < y0:
        y0, y1 = y1, y0

    if shape.fill_rgb is not None:
        draw.rectangle([x0, y0, x1, y1], fill=shape.fill_rgb + (255,))

    if shape.stroke_rgb is not None:
        sw = max(1, int(round(_emu_to_px_f(shape.stroke_width_emu or 9525))))
        draw.rectangle(
            [x0, y0, x1, y1],
            outline=shape.stroke_rgb + (255,),
            width=sw,
        )

    if shape.text_runs:
        _draw_text_in_box(canvas, draw, shape, x0, y0, x1, y1)


def _draw_ellipse_shape(draw, canvas, shape: FloatingShape,
                        px_x, px_y) -> None:
    x0 = px_x(shape.x_emu)
    y0 = px_y(shape.y_emu)
    x1 = px_x(shape.x_emu + shape.width_emu)
    y1 = px_y(shape.y_emu + shape.height_emu)
    if x1 < x0:
        x0, x1 = x1, x0
    if y1 < y0:
        y0, y1 = y1, y0

    if shape.fill_rgb is not None:
        draw.ellipse([x0, y0, x1, y1], fill=shape.fill_rgb + (255,))
    if shape.stroke_rgb is not None:
        sw = max(1, int(round(_emu_to_px_f(shape.stroke_width_emu or 9525))))
        draw.ellipse(
            [x0, y0, x1, y1],
            outline=shape.stroke_rgb + (255,),
            width=sw,
        )

    if shape.text_runs:
        _draw_text_in_box(canvas, draw, shape, x0, y0, x1, y1)


def _draw_line_shape(draw, shape: FloatingShape, px_x, px_y) -> None:
    # The line spans the bounding box diagonally; flipH / flipV choose
    # which diagonal. With no flips, line runs from top-left to bottom-right.
    x_lo = px_x(shape.x_emu)
    y_lo = px_y(shape.y_emu)
    x_hi = px_x(shape.x_emu + shape.width_emu)
    y_hi = px_y(shape.y_emu + shape.height_emu)

    if shape.flip_h:
        x_start, x_end = x_hi, x_lo
    else:
        x_start, x_end = x_lo, x_hi
    if shape.flip_v:
        y_start, y_end = y_hi, y_lo
    else:
        y_start, y_end = y_lo, y_hi

    color_rgb = shape.stroke_rgb if shape.stroke_rgb is not None else (0, 0, 0)
    color = color_rgb + (255,)
    sw = max(1, int(round(_emu_to_px_f(shape.stroke_width_emu or 9525))))

    draw.line(
        [(x_start, y_start), (x_end, y_end)],
        fill=color,
        width=sw,
    )

    # Arrowheads. headEnd is at (x_start, y_start) per OOXML convention
    # (head = where the line "starts" in shape coordinates); tailEnd is at
    # (x_end, y_end). We always draw at the actual endpoint, with the
    # triangle pointing along the line direction towards that endpoint.
    if shape.head_end_present:
        _draw_arrowhead(draw, x_end, y_end, x_start, y_start, color, sw)
    if shape.tail_end_present:
        _draw_arrowhead(draw, x_start, y_start, x_end, y_end, color, sw)


def _draw_arrowhead(draw, from_x: float, from_y: float,
                    to_x: float, to_y: float,
                    fill, stroke_width: int) -> None:
    """Draw a filled triangular arrowhead pointing from (from_x, from_y)
    toward (to_x, to_y), with the tip at (to_x, to_y)."""
    import math
    dx = to_x - from_x
    dy = to_y - from_y
    length = math.hypot(dx, dy)
    if length < 1e-6:
        return
    ux, uy = dx / length, dy / length
    # Arrowhead size scales with stroke width; Word uses several preset
    # sizes but for v1 we pick a single sensible default.
    head_len = max(6.0, stroke_width * 4.0)
    head_half_w = max(4.0, stroke_width * 2.5)
    # Base centre (back of the triangle).
    bx = to_x - ux * head_len
    by = to_y - uy * head_len
    # Perpendicular unit vector.
    px, py = -uy, ux
    p1 = (to_x, to_y)
    p2 = (bx + px * head_half_w, by + py * head_half_w)
    p3 = (bx - px * head_half_w, by - py * head_half_w)
    draw.polygon([p1, p2, p3], fill=fill)


def _draw_text_in_box(canvas, draw, shape: FloatingShape,
                      x0: float, y0: float, x1: float, y1: float) -> None:
    """Render text_runs inside the rectangle [x0,y0,x1,y1]. Applies
    bodyPr insets to derive the content area, then uses bodyPr anchor to
    position the text block vertically (t/ctr/b). Lines are centred
    horizontally within the content area. v1 limitations:
    - no word wrapping (text overflows if too long)
    - single font per shape (uses the first run's size)
    - colour from the first run (or black)
    - horizontal alignment always centred (Word's default for shapes)"""
    if not shape.text_runs:
        return

    # Convert EMU insets to pixels using the box's pixel-per-EMU scale.
    # We derive scale from the shape's bounding box because that is the only
    # information we have linking EMU to the rendered pixels for this shape.
    box_w_emu = max(1, shape.width_emu)
    box_h_emu = max(1, shape.height_emu)
    box_w_px = x1 - x0
    box_h_px = y1 - y0
    sx = box_w_px / box_w_emu
    sy = box_h_px / box_h_emu

    inset_l = shape.text_ins_l * sx
    inset_r = shape.text_ins_r * sx
    inset_t = shape.text_ins_t * sy
    inset_b = shape.text_ins_b * sy

    content_x0 = x0 + inset_l
    content_x1 = x1 - inset_r
    content_y0 = y0 + inset_t
    content_y1 = y1 - inset_b
    if content_x1 <= content_x0:
        content_x0, content_x1 = x0, x1
    if content_y1 <= content_y0:
        content_y0, content_y1 = y0, y1
    content_w = content_x1 - content_x0
    content_h = content_y1 - content_y0

    # Default styling from the first run.
    first_fmt = shape.text_runs[0][1]
    size_pt = first_fmt.get("size_pt") or 11.0
    size_px = int(round(size_pt * IMAGE_DPI / 72.0))
    color_rgb = first_fmt.get("rgb") or (0, 0, 0)
    bold = first_fmt.get("bold", False)
    italic = first_fmt.get("italic", False)

    # Concatenate runs into lines (split on newline tokens).
    text = "".join(t for t, _ in shape.text_runs)
    lines = text.splitlines() or [text]

    font = _pick_font_styled(size_px, bold=bold, italic=italic)

    # Measure line heights with the picked font.
    try:
        line_heights = [
            (draw.textbbox((0, 0), line, font=font)[3]
             - draw.textbbox((0, 0), line, font=font)[1])
            or size_px
            for line in lines
        ]
    except Exception:
        line_heights = [size_px] * len(lines)
    total_h = sum(line_heights)

    # Vertical anchor: t = top, ctr = centre, b = bottom.
    if shape.text_anchor == "ctr":
        cur_y = content_y0 + max(0, (content_h - total_h) / 2.0)
    elif shape.text_anchor == "b":
        cur_y = content_y1 - total_h
    else:
        cur_y = content_y0

    for line, lh in zip(lines, line_heights):
        try:
            bbox = draw.textbbox((0, 0), line, font=font)
            line_w = bbox[2] - bbox[0]
        except Exception:
            line_w = size_px * len(line)
        tx = content_x0 + max(0, (content_w - line_w) / 2.0)
        draw.text((tx, cur_y), line, fill=color_rgb + (255,), font=font)
        cur_y += lh


def _pick_font_styled(size_px: int, *, bold: bool, italic: bool):
    """Pick a font, attempting bold/italic variants on Windows."""
    from PIL import ImageFont
    base_candidates = []
    if bold and italic:
        base_candidates = [
            "C:\\Windows\\Fonts\\arialbi.ttf",
            "C:\\Windows\\Fonts\\calibriz.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
        ]
    elif bold:
        base_candidates = [
            "C:\\Windows\\Fonts\\arialbd.ttf",
            "C:\\Windows\\Fonts\\calibrib.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ]
    elif italic:
        base_candidates = [
            "C:\\Windows\\Fonts\\ariali.ttf",
            "C:\\Windows\\Fonts\\calibrii.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
        ]
    for path in base_candidates:
        try:
            return ImageFont.truetype(path, size=max(8, int(size_px)))
        except (OSError, IOError):
            continue
    return _pick_font(size_px)


# ---------------------------------------------------------------------------
# Tracked changes
# ---------------------------------------------------------------------------

def _check_tracked_changes(doc, docx_path: Path,
                           tracker: PageTracker) -> None:
    body_xml = doc.element.body
    ins = body_xml.find(f".//{{{NS['w']}}}ins")
    dele = body_xml.find(f".//{{{NS['w']}}}del")
    offender = ins if ins is not None else dele
    if offender is not None:
        page = tracker.page_for(offender)
        _fail_with_page(
            f"{docx_path.name} contains unresolved tracked changes. "
            f"Accept or reject all changes in Word before running --togit.",
            page,
        )


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

class ExtractedImage:
    __slots__ = ("rendered_bytes", "ext", "hash_hex", "alt_text",
                 "width_px", "height_px")

    def __init__(self, rendered_bytes: bytes, ext: str, alt_text: str,
                 width_px: int, height_px: int) -> None:
        self.rendered_bytes = rendered_bytes
        self.ext = ext
        self.alt_text = alt_text
        self.width_px = width_px
        self.height_px = height_px
        self.hash_hex = hashlib.sha256(rendered_bytes).hexdigest()


def _emu_to_px(emu: int, dpi: int = IMAGE_DPI) -> int:
    # 914400 EMU per inch.
    return max(1, int(round(emu * dpi / 914400)))


def _extract_inline_image(drawing_el, doc, docx_path: Path,
                          page: int,
                          overlay_shapes: List[FloatingShape]
                          ) -> ExtractedImage:
    """Render a single inline image found inside a <w:drawing> element,
    optionally compositing floating annotation shapes on top."""
    # Reject anchored/floating images.
    if drawing_el.find(f"{{{NS['wp']}}}anchor") is not None:
        _fail_with_page(
            f"{docx_path.name} contains a floating or anchored image. "
            f"Only inline images are supported.",
            page,
        )

    inline = drawing_el.find(f"{{{NS['wp']}}}inline")
    if inline is None:
        _fail_with_page(
            f"{docx_path.name} contains a drawing element with no inline "
            f"layout; only inline images are supported.",
            page,
        )

    # Extent (display size in EMU).
    extent = inline.find(f"{{{NS['wp']}}}extent")
    if extent is None:
        _fail_with_page(
            f"{docx_path.name} has an inline image with no display extent.",
            page,
        )
    try:
        cx = int(extent.get("cx", "0"))
        cy = int(extent.get("cy", "0"))
    except ValueError:
        _fail_with_page(
            f"{docx_path.name} has an inline image with non-numeric extent.",
            page,
        )
        cx = cy = 0  # unreachable

    # Alt text from docPr/@descr (preferred) or @title.
    doc_pr = inline.find(f"{{{NS['wp']}}}docPr")
    alt_text = ""
    if doc_pr is not None:
        alt_text = (doc_pr.get("descr") or doc_pr.get("title") or "").strip()

    # Find the blip (image reference) and optional srcRect (cropping).
    blip = inline.find(f".//{{{NS['a']}}}blip")
    if blip is None:
        _fail_with_page(
            f"{docx_path.name} has an inline image with no blip reference.",
            page,
        )
    r_embed = blip.get(f"{{{NS['r']}}}embed")
    if not r_embed:
        _fail_with_page(
            f"{docx_path.name} has an inline image with no r:embed "
            f"relationship.",
            page,
        )

    src_rect = inline.find(f".//{{{NS['a']}}}srcRect")

    # Resolve relationship to image part.
    try:
        image_part = doc.part.related_parts[r_embed]
    except KeyError:
        _fail_with_page(
            f"{docx_path.name} references missing image relationship "
            f"{r_embed!r}.",
            page,
        )
        return None  # unreachable

    content_type = image_part.content_type
    if content_type not in SUPPORTED_IMAGE_CONTENT_TYPES:
        _fail_with_page(
            f"{docx_path.name} contains an unsupported image type "
            f"{content_type!r}. Only PNG and JPEG are supported.",
            page,
        )
        return None  # unreachable
    out_ext = SUPPORTED_IMAGE_CONTENT_TYPES[content_type]

    # Render via Pillow.
    try:
        src_img = Image.open(io.BytesIO(image_part.blob))
        src_img.load()
    except Exception as exc:
        _fail_with_page(
            f"Failed to decode an image in {docx_path.name}: {exc}",
            page,
        )
        return None  # unreachable

    # Apply cropping if specified. srcRect values are in 1/100000 (i.e.
    # thousandths of a percent). Positive values trim from that edge.
    if src_rect is not None:
        try:
            l = int(src_rect.get("l", "0")) / 100000.0
            t = int(src_rect.get("t", "0")) / 100000.0
            r = int(src_rect.get("r", "0")) / 100000.0
            b = int(src_rect.get("b", "0")) / 100000.0
        except ValueError:
            _fail_with_page(
                f"{docx_path.name} has non-numeric image crop metadata.",
                page,
            )
            return None  # unreachable
        w, h = src_img.size
        left = max(0, int(round(w * l)))
        top = max(0, int(round(h * t)))
        right = min(w, int(round(w * (1.0 - r))))
        bottom = min(h, int(round(h * (1.0 - b))))
        if right <= left or bottom <= top:
            _fail_with_page(
                f"{docx_path.name} has an image crop that collapses to "
                f"zero area; cannot render reliably.",
                page,
            )
            return None  # unreachable
        src_img = src_img.crop((left, top, right, bottom))

    # Resize to Word's display size. Always rescale (up or down) so images
    # rendered at the same width in Word produce assets at the same width.
    # Aspect ratio comes from Word's (cx, cy) extent, which mirrors the
    # source proportions unless the user deliberately distorted them.
    target_w = _emu_to_px(cx)
    target_h = _emu_to_px(cy)
    if (src_img.width, src_img.height) != (target_w, target_h):
        src_img = src_img.resize((target_w, target_h), Image.LANCZOS)

    # Composite floating annotation shapes onto the rendered image.
    if overlay_shapes:
        src_img = _composite_shapes(src_img, overlay_shapes, (cx, cy))
        # Compositing produces RGBA; ensure the output format can handle it.
        if out_ext == "jpg" and src_img.mode == "RGBA":
            # JPEG has no alpha channel; flatten onto white before save.
            from PIL import Image as PILImage
            bg = PILImage.new("RGB", src_img.size, (255, 255, 255))
            bg.paste(src_img, mask=src_img.split()[3])
            src_img = bg

    # Encode bytes deterministically.
    buf = io.BytesIO()
    try:
        if out_ext == "png":
            save_img = src_img
            if save_img.mode == "P":
                save_img = save_img.convert("RGBA")
            save_img.save(buf, format="PNG", optimize=True)
        else:
            save_img = src_img.convert("RGB") if src_img.mode != "RGB" else src_img
            save_img.save(buf, format="JPEG", quality=90, optimize=True)
    except Exception as exc:
        _fail_with_page(
            f"Failed to render an image in {docx_path.name}: {exc}",
            page,
        )
        return None  # unreachable

    return ExtractedImage(
        rendered_bytes=buf.getvalue(),
        ext=out_ext,
        alt_text=alt_text,
        width_px=src_img.width,
        height_px=src_img.height,
    )


# ---------------------------------------------------------------------------
# Markdown emission helpers
# ---------------------------------------------------------------------------

def _escape_md_text(text: str, *, in_table: bool = False) -> str:
    """Escape markdown-significant characters in literal text.

    Conservative but not paranoid; deliberately does not escape every
    character, only those likely to be misinterpreted by a GFM parser.
    """
    if not text:
        return text
    # Escape backslash first to avoid double-escaping.
    out = text.replace("\\", "\\\\")
    # Escape backticks, brackets, asterisks, underscores, pipes (in tables).
    for ch in "`*_[]<>":
        out = out.replace(ch, "\\" + ch)
    if in_table:
        out = out.replace("|", "\\|")
        # Newlines in a cell would break the row; collapse to spaces.
        out = out.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    return out


# ---------------------------------------------------------------------------
# DOCX -> Markdown converter
# ---------------------------------------------------------------------------

class _RunPiece:
    """One contiguous text piece with formatting flags and optional link."""

    __slots__ = ("text", "bold", "italic", "code", "strike", "link",
                 "is_image", "image_index")

    def __init__(self, text: str = "", *, bold: bool = False,
                 italic: bool = False, code: bool = False,
                 strike: bool = False, link: Optional[str] = None,
                 is_image: bool = False,
                 image_index: Optional[int] = None) -> None:
        self.text = text
        self.bold = bold
        self.italic = italic
        self.code = code
        self.strike = strike
        self.link = link
        self.is_image = is_image
        self.image_index = image_index

    def fmt_key(self) -> Tuple:
        return (self.bold, self.italic, self.code, self.strike, self.link)


class _NumberingResolver:
    """Resolves (numId, ilvl) -> 'bullet' or 'ordered' by reading the
    document's word/numbering.xml part.

    Word's UI ribbon buttons for numbered/bulleted lists do NOT apply the
    'List Number'/'List Bullet' styles; instead they apply 'List Paragraph'
    style with a <w:numPr> reference to a numbering definition. Detecting
    ordered vs bullet for those paragraphs requires reading numbering.xml
    and looking up the level's <w:numFmt>.

    Anything with numFmt 'bullet' is unordered; everything else (decimal,
    lowerLetter/upperLetter, lowerRoman/upperRoman, decimalZero, etc.) is
    treated as ordered. Style-link references (<w:numStyleLink>,
    <w:styleLink>) are not resolved in v1; the lookup returns None for
    those cases and the caller falls back to a conservative default.
    """

    def __init__(self, doc) -> None:
        self._kind: Dict[Tuple[str, str], str] = {}
        self._load(doc)

    def _load(self, doc) -> None:
        try:
            numbering_part = doc.part.numbering_part
        except Exception:
            numbering_part = None
        if numbering_part is None:
            return
        root = getattr(numbering_part, "element", None)
        if root is None:
            return

        # abstractNumId -> {ilvl: numFmt}
        abstract: Dict[str, Dict[str, str]] = {}
        for an in root.findall(qn("w:abstractNum")):
            an_id = an.get(qn("w:abstractNumId"))
            if an_id is None:
                continue
            levels: Dict[str, str] = {}
            for lvl in an.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                nf = lvl.find(qn("w:numFmt"))
                if ilvl is not None and nf is not None:
                    levels[ilvl] = nf.get(qn("w:val"), "")
            abstract[an_id] = levels

        # numId -> abstractNumId (+ per-level overrides via w:lvlOverride)
        for num in root.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            if num_id is None:
                continue
            abs_ref = num.find(qn("w:abstractNumId"))
            if abs_ref is None:
                continue
            an_id = abs_ref.get(qn("w:val"))
            if an_id is None:
                continue
            levels = dict(abstract.get(an_id, {}))

            for ov in num.findall(qn("w:lvlOverride")):
                ov_ilvl = ov.get(qn("w:ilvl"))
                ov_lvl = ov.find(qn("w:lvl"))
                if ov_ilvl is not None and ov_lvl is not None:
                    nf = ov_lvl.find(qn("w:numFmt"))
                    if nf is not None:
                        levels[ov_ilvl] = nf.get(qn("w:val"), "")

            for ilvl, fmt in levels.items():
                kind = "bullet" if fmt == "bullet" else "ordered"
                self._kind[(num_id, ilvl)] = kind

    def kind_of(self, num_id: Optional[str],
                ilvl: Optional[str]) -> Optional[str]:
        """Returns 'bullet', 'ordered', or None if the lookup could not be
        resolved (e.g. unknown numId, style-link reference, no numbering
        part)."""
        if num_id is None:
            return None
        return self._kind.get((num_id, ilvl or "0"))


class DocxToMarkdown:
    def __init__(self, doc, docx_path: Path,
                 image_filenames: List[str],
                 tracker: Optional[PageTracker] = None) -> None:
        self.doc = doc
        self.docx_path = docx_path
        # image_filenames[i] is the final filename (no path) for the i-th
        # inline image encountered in document order.
        self.image_filenames = image_filenames
        self._image_counter = 0  # next image index to consume
        self.tracker = tracker
        self._numbering = _NumberingResolver(doc)

    # --- public ---

    def render(self) -> str:
        blocks: List[str] = []
        body = self.doc.element.body
        items = list(body.iterchildren())

        i = 0
        while i < len(items):
            el = items[i]
            tag = el.tag

            if tag == qn("w:p"):
                para = DocxParagraph(el, self.doc)
                style_name = (para.style.name if para.style else "Normal") or "Normal"

                # Group consecutive CodeBlock paragraphs into one fenced block.
                if style_name == CODE_BLOCK_STYLE_NAME:
                    code_lines: List[str] = []
                    while i < len(items) and items[i].tag == qn("w:p"):
                        p2 = DocxParagraph(items[i], self.doc)
                        s2 = (p2.style.name if p2.style else "Normal") or "Normal"
                        if s2 != CODE_BLOCK_STYLE_NAME:
                            break
                        code_lines.append(self._raw_text(p2))
                        i += 1
                    blocks.append("```\n" + "\n".join(code_lines) + "\n```")
                    continue

                # Group consecutive Quote paragraphs into one blockquote.
                if style_name == QUOTE_STYLE_NAME:
                    quote_paras: List[str] = []
                    while i < len(items) and items[i].tag == qn("w:p"):
                        p2 = DocxParagraph(items[i], self.doc)
                        s2 = (p2.style.name if p2.style else "Normal") or "Normal"
                        if s2 != QUOTE_STYLE_NAME:
                            break
                        rendered = self._render_inline_paragraph(p2).strip()
                        quote_paras.append(rendered)
                        i += 1
                    # Emit: each paragraph prefixed with "> "; insert a bare
                    # ">" line between paragraphs so GFM treats them as
                    # separate paragraphs inside the blockquote.
                    quote_lines: List[str] = []
                    for idx_q, qp in enumerate(quote_paras):
                        if idx_q > 0:
                            quote_lines.append(">")
                        if qp == "":
                            quote_lines.append(">")
                        else:
                            for sub in qp.split("\n"):
                                quote_lines.append(
                                    ("> " + sub) if sub else ">"
                                )
                    blocks.append("\n".join(quote_lines))
                    continue

                # Group consecutive list paragraphs into one list block.
                if self._is_list_paragraph(para):
                    list_lines, consumed = self._render_list(items, i)
                    blocks.append("\n".join(list_lines))
                    i += consumed
                    continue

                # Headings.
                m = re.match(r"^Heading (\d)$", style_name)
                if m:
                    level = int(m.group(1))
                    inline = self._render_inline_paragraph(para)
                    blocks.append(("#" * level) + " " + inline.strip())
                    i += 1
                    continue

                # Horizontal rule: empty paragraph with bottom border.
                if self._is_hr(para):
                    blocks.append("---")
                    i += 1
                    continue

                # Default: normal paragraph.
                inline = self._render_inline_paragraph(para)
                if inline.strip() == "":
                    # Skip stray empty paragraphs to avoid double blank lines.
                    i += 1
                    continue
                blocks.append(inline)
                i += 1

            elif tag == qn("w:tbl"):
                table = DocxTable(el, self.doc)
                blocks.append(self._render_table(table))
                i += 1

            else:
                # Section properties, bookmarks etc - ignored.
                i += 1

        # All inline images should have been consumed.
        if self._image_counter != len(self.image_filenames):
            fail(
                f"{self.docx_path.name}: rendered {self._image_counter} of "
                f"{len(self.image_filenames)} inline images. Document "
                f"structure may contain unsupported image placement."
            )

        # Join with blank lines between blocks. Trailing newline.
        return "\n\n".join(blocks).rstrip() + "\n"

    # --- list rendering ---

    def _is_list_paragraph(self, para: DocxParagraph) -> bool:
        style_name = (para.style.name if para.style else "") or ""
        if style_name in ("List Bullet", "List Number"):
            return True
        # User-added Word lists may use numPr without our style names.
        numPr = self._numpr(para)
        if numPr is None:
            # "List Paragraph" without numPr is not a list (just a styled
            # paragraph leftover from a removed list).
            return False
        # numId=0 means "no list reference" (explicit removal from a list).
        numId_el = numPr.find(qn("w:numId"))
        if numId_el is None:
            return False
        if numId_el.get(qn("w:val"), "0") == "0":
            return False
        return True

    def _numpr(self, para: DocxParagraph):
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        return pPr.find(qn("w:numPr"))

    def _list_level(self, para: DocxParagraph) -> int:
        numPr = self._numpr(para)
        if numPr is not None:
            ilvl_el = numPr.find(qn("w:ilvl"))
            if ilvl_el is not None:
                try:
                    return max(0, int(ilvl_el.get(qn("w:val"), "0")))
                except ValueError:
                    pass
        # Fall back to left_indent (set by --fromgit for nested levels).
        indent = para.paragraph_format.left_indent
        if indent is None:
            return 0
        # indent is in EMU (twentieths of a point); python-docx returns Emu.
        try:
            indent_cm = indent.cm
        except AttributeError:
            indent_cm = float(indent) / 360000.0
        if indent_cm <= LIST_INDENT_STEP_CM + 0.01:
            return 0
        return max(0, int(round((indent_cm - LIST_INDENT_STEP_CM) /
                                LIST_INDENT_STEP_CM)))

    def _is_ordered_list(self, para: DocxParagraph) -> bool:
        # Explicit style names take precedence (these are what --fromgit
        # writes) and avoid any dependency on numbering.xml.
        style_name = (para.style.name if para.style else "") or ""
        if style_name == "List Number":
            return True
        if style_name == "List Bullet":
            return False

        # Word-UI-created lists use numPr referring to a definition in
        # numbering.xml. Look up the numFmt for this (numId, ilvl).
        numPr = self._numpr(para)
        if numPr is None:
            return False
        numId_el = numPr.find(qn("w:numId"))
        if numId_el is None:
            return False
        num_id = numId_el.get(qn("w:val"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
        kind = self._numbering.kind_of(num_id, ilvl)
        if kind == "ordered":
            return True
        if kind == "bullet":
            return False
        # Unresolved (e.g. style-link reference, or numbering.xml missing).
        # Conservative default: bullet.
        return False

    def _render_list(self, items: List, start: int) -> Tuple[List[str], int]:
        lines: List[str] = []
        # Track ordering counter per level so ordered list markers count up.
        counters: Dict[int, int] = {}
        last_level = -1
        consumed = 0
        # Each ordered list at each level should restart at 1 when we move
        # to a new list. We reset a level's counter when we leave it (move
        # to a shallower level) or when its parent type changes.

        i = start
        while i < len(items) and items[i].tag == qn("w:p"):
            para = DocxParagraph(items[i], self.doc)
            if not self._is_list_paragraph(para):
                break
            level = self._list_level(para)
            ordered = self._is_ordered_list(para)
            # Reset counters for any level deeper than current (we're moving
            # back up the tree).
            for deeper in [lvl for lvl in counters if lvl > level]:
                del counters[deeper]
            if ordered:
                counters[level] = counters.get(level, 0) + 1
                marker = f"{counters[level]}. "
            else:
                # Bullets do not need counting; clear any prior ordered count
                # at this level so a subsequent ordered list restarts at 1.
                counters.pop(level, None)
                marker = "- "
            indent = "  " * level
            content = self._render_inline_paragraph(para).strip()
            # Detect task list pattern. After inline rendering, brackets have
            # been escaped to "\[" and "\]"; rewrite the leading marker so it
            # round-trips as a real GFM task-list item.
            task_match = re.match(
                r"^\\\[(?P<state>[ xX])\\\] ?(?P<rest>.*)$",
                content,
            )
            if task_match:
                state = task_match.group("state").lower()
                if state == " ":
                    state = " "
                content = f"[{state}] {task_match.group('rest')}"
            lines.append(f"{indent}{marker}{content}")
            i += 1
            consumed += 1
            last_level = level
        return lines, consumed

    # --- table rendering ---

    def _render_table(self, table: DocxTable) -> str:
        # Validate: no merged cells, no nested tables, no images in cells.
        rows = table.rows
        if not rows:
            return ""
        page = self.tracker.page_for(table._tbl) if self.tracker else None
        for row in rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is not None:
                    if tcPr.find(qn("w:gridSpan")) is not None:
                        _fail_with_page(
                            f"{self.docx_path.name} contains a table with "
                            f"horizontally merged cells (gridSpan), which "
                            f"GFM pipe tables cannot represent.",
                            page,
                        )
                    if tcPr.find(qn("w:vMerge")) is not None:
                        _fail_with_page(
                            f"{self.docx_path.name} contains a table with "
                            f"vertically merged cells (vMerge), which GFM "
                            f"pipe tables cannot represent.",
                            page,
                        )
                if tc.find(f".//{{{NS['w']}}}tbl") is not None:
                    _fail_with_page(
                        f"{self.docx_path.name} contains a nested table, "
                        f"which is unsupported.",
                        page,
                    )
                if tc.find(f".//{{{NS['w']}}}drawing") is not None:
                    _fail_with_page(
                        f"{self.docx_path.name} contains an image inside a "
                        f"table cell, which is unsupported.",
                        page,
                    )

        ncols = len(rows[0].cells)
        # Detect column count mismatch across rows: also unsupported.
        for r in rows[1:]:
            if len(r.cells) != ncols:
                _fail_with_page(
                    f"{self.docx_path.name} contains a table with ragged "
                    f"row widths; cannot map to a GFM pipe table.",
                    page,
                )

        # First row is the header. Build cell text and per-column alignment.
        header_cells: List[str] = []
        col_aligns: List[str] = []
        for cell in rows[0].cells:
            header_cells.append(self._render_cell(cell, in_header=True))
            align = self._cell_alignment(cell)
            col_aligns.append(align)

        body_rows: List[List[str]] = []
        for row in rows[1:]:
            body_rows.append([self._render_cell(c) for c in row.cells])

        # Compose pipe table.
        out = ["| " + " | ".join(header_cells) + " |"]
        sep_parts = []
        for align in col_aligns:
            if align == "center":
                sep_parts.append(":---:")
            elif align == "right":
                sep_parts.append("---:")
            elif align == "left":
                sep_parts.append(":---")
            else:
                sep_parts.append("---")
        out.append("| " + " | ".join(sep_parts) + " |")
        for body in body_rows:
            out.append("| " + " | ".join(body) + " |")
        return "\n".join(out)

    def _render_cell(self, cell, *, in_header: bool = False) -> str:
        pieces: List[str] = []
        for p in cell.paragraphs:
            inline = self._render_inline_paragraph(
                p, in_table=True, ignore_bold=in_header,
            )
            if inline.strip():
                pieces.append(inline.strip())
        return " ".join(pieces) if pieces else ""

    def _cell_alignment(self, cell) -> Optional[str]:
        for p in cell.paragraphs:
            jc = p._p.find(f".//{qn('w:jc')}")
            if jc is not None:
                val = jc.get(qn("w:val"), "")
                if val in ("center", "right", "left"):
                    return val
            if p.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                return "center"
            if p.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                return "right"
            if p.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                return "left"
        return None

    # --- HR and miscellaneous ---

    def _is_hr(self, para: DocxParagraph) -> bool:
        if para.text.strip() != "":
            return False
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            return False
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return False
        return pBdr.find(qn("w:bottom")) is not None

    def _raw_text(self, para: DocxParagraph) -> str:
        """Return paragraph text verbatim (used for code blocks)."""
        # Walk all w:t and w:tab/w:br nodes in order.
        parts: List[str] = []
        for el in para._p.iter():
            t = el.tag
            if t == qn("w:t"):
                parts.append(el.text or "")
            elif t == qn("w:tab"):
                parts.append("\t")
            elif t == qn("w:br"):
                parts.append("\n")
        return "".join(parts)

    # --- inline paragraph rendering (bold/italic/code/links/images) ---

    def _render_inline_paragraph(self, para: DocxParagraph, *,
                                 in_table: bool = False,
                                 ignore_bold: bool = False) -> str:
        pieces = self._collect_pieces(para)
        if ignore_bold:
            for piece in pieces:
                piece.bold = False
        return self._emit_pieces(pieces, in_table=in_table)

    def _collect_pieces(self, para: DocxParagraph) -> List[_RunPiece]:
        """Walk paragraph XML and produce a flat list of run pieces."""
        pieces: List[_RunPiece] = []

        def walk(node, link_url: Optional[str]) -> None:
            for child in node:
                tag = child.tag
                if tag == qn("w:hyperlink"):
                    # Resolve r:id to URL via paragraph's part rels.
                    r_id = child.get(qn("r:id"))
                    url = None
                    if r_id:
                        try:
                            url = para.part.rels[r_id].target_ref
                        except KeyError:
                            url = None
                    walk(child, url)
                elif tag == qn("w:r"):
                    self._process_run(child, pieces, link_url)
                elif tag == qn("w:smartTag") or tag == qn("w:ins"):
                    # Tracked changes already rejected earlier; smartTag
                    # is a wrapper around runs.
                    walk(child, link_url)
                # other children (bookmarks, comment refs, etc) ignored

        walk(para._p, None)
        return pieces

    def _process_run(self, r_el, pieces: List[_RunPiece],
                     link_url: Optional[str]) -> None:
        rPr = r_el.find(qn("w:rPr"))
        bold = False
        italic = False
        code = False
        strike = False
        if rPr is not None:
            # Style reference for inline code.
            rStyle = rPr.find(qn("w:rStyle"))
            if rStyle is not None:
                if rStyle.get(qn("w:val"), "") == INLINE_CODE_STYLE_NAME:
                    code = True
            # Bold: <w:b/> or <w:b w:val="true"/>; <w:b w:val="false"/> means off.
            b = rPr.find(qn("w:b"))
            if b is not None and b.get(qn("w:val"), "true").lower() not in ("0", "false"):
                bold = True
            i = rPr.find(qn("w:i"))
            if i is not None and i.get(qn("w:val"), "true").lower() not in ("0", "false"):
                italic = True
            s = rPr.find(qn("w:strike"))
            if s is not None and s.get(qn("w:val"), "true").lower() not in ("0", "false"):
                strike = True

        # Walk run children in order for text, breaks, drawings.
        for child in r_el:
            tag = child.tag
            if tag == qn("w:t"):
                if child.text:
                    pieces.append(_RunPiece(
                        text=child.text,
                        bold=bold, italic=italic, code=code, strike=strike,
                        link=link_url,
                    ))
            elif tag == qn("w:tab"):
                pieces.append(_RunPiece(
                    text="\t",
                    bold=bold, italic=italic, code=code, strike=strike,
                    link=link_url,
                ))
            elif tag == qn("w:br"):
                # Markdown line break: two trailing spaces + newline. But
                # for v1 we collapse to a space to keep paragraph flow
                # readable; explicit hard breaks remain rare in practice.
                pieces.append(_RunPiece(
                    text="  \n",
                    bold=False, italic=False, code=False, strike=False,
                    link=None,
                ))
            elif tag == qn("w:drawing"):
                # Only INLINE drawings consume an image slot. Anchored
                # drawings (annotation shapes) have already been composited
                # into their host inline image by the extractor.
                if child.find(f"{{{NS['wp']}}}inline") is None:
                    continue
                if self._image_counter >= len(self.image_filenames):
                    fail(
                        f"{self.docx_path.name}: encountered more inline "
                        f"images than were extracted."
                    )
                idx = self._image_counter
                pieces.append(_RunPiece(is_image=True, image_index=idx,
                                        link=link_url))
                self._image_counter += 1

            elif tag == f"{{{NS['mc']}}}AlternateContent":
                # Inline images may be wrapped in mc:AlternateContent
                # (modern format with VML fallback). Look for an inline
                # drawing inside the mc:Choice subtree.
                choice = child.find(f"{{{NS['mc']}}}Choice")
                if choice is None:
                    continue
                for d in _iter_drawings_in(choice):
                    if d.find(f"{{{NS['wp']}}}inline") is None:
                        continue
                    if self._image_counter >= len(self.image_filenames):
                        fail(
                            f"{self.docx_path.name}: encountered more "
                            f"inline images than were extracted."
                        )
                    idx = self._image_counter
                    pieces.append(_RunPiece(is_image=True, image_index=idx,
                                            link=link_url))
                    self._image_counter += 1

    def _emit_pieces(self, pieces: List[_RunPiece], *,
                     in_table: bool) -> str:
        # First, coalesce adjacent same-format text pieces (excluding images).
        out: List[str] = []
        i = 0
        while i < len(pieces):
            p = pieces[i]
            if p.is_image:
                out.append(self._format_image(p))
                i += 1
                continue
            # Hard break tokens are preserved verbatim.
            if p.text == "  \n":
                out.append(p.text)
                i += 1
                continue
            # Coalesce same-format run.
            j = i
            buf = []
            key = p.fmt_key()
            while j < len(pieces):
                q = pieces[j]
                if q.is_image or q.text == "  \n" or q.fmt_key() != key:
                    break
                buf.append(q.text)
                j += 1
            text = "".join(buf)
            out.append(self._format_text_run(text, key, in_table=in_table))
            i = j
        return "".join(out)

    def _format_text_run(self, text: str, fmt_key: Tuple,
                         *, in_table: bool) -> str:
        bold, italic, code, strike, link = fmt_key
        if not text:
            return ""
        if code:
            # Inline code: escape backticks by widening fence.
            fence = "`"
            while fence in text:
                fence += "`"
            # Pad if text starts/ends with backtick.
            pad = " " if text.startswith("`") or text.endswith("`") else ""
            rendered = f"{fence}{pad}{text}{pad}{fence}"
        else:
            rendered = _escape_md_text(text, in_table=in_table)
            if bold and italic:
                rendered = f"***{rendered}***"
            elif bold:
                rendered = f"**{rendered}**"
            elif italic:
                rendered = f"*{rendered}*"
            if strike:
                rendered = f"~~{rendered}~~"
        if link:
            # Hyperlink wraps the rendered content.
            rendered = f"[{rendered}]({link})"
        return rendered

    def _format_image(self, piece: _RunPiece) -> str:
        from urllib.parse import quote
        idx = piece.image_index
        filename = self.image_filenames[idx]
        alt = self._image_alts[idx] if idx < len(self._image_alts) else ""
        # URL-encode the path so spaces and other special characters survive
        # rendering. Keep "/" unencoded so the assets directory separator
        # remains readable.
        rel = f"{ASSETS_DIR_NAME}/{quote(filename, safe='-_.')}"
        md = f"![{_escape_md_text(alt)}]({rel})"
        if piece.link:
            md = f"[{md}]({piece.link})"
        return md


# ---------------------------------------------------------------------------
# Conversion orchestration
# ---------------------------------------------------------------------------

def _collect_inline_images_with_overlays(doc, docx_path: Path,
                                         tracker: PageTracker):
    """Walk the document body and yield (drawing_el, page, overlay_shapes)
    tuples in document order for every inline image.

    Each paragraph's anchored shapes are gathered, then assigned to the
    inline image they sit on top of (containment-tested against that
    image's extent). Anchored shapes that are NOT wholly inside an inline
    image, or paragraphs that have multiple inline images with anchored
    shapes (ambiguous attribution), trigger fail-fast with the page
    number."""
    body = doc.element.body
    for p_el in body.iter(qn("w:p")):
        # Discover this paragraph's drawings, excluding mc:Fallback copies.
        inline_drawings: List = []
        anchored_drawings: List = []
        for d in _iter_drawings_in(p_el):
            # An ancestor walk to confirm the drawing belongs to THIS
            # paragraph and not a nested paragraph (e.g. inside a textbox
            # within a shape on this paragraph).
            ancestor = d.getparent()
            while ancestor is not None and ancestor is not p_el:
                if ancestor.tag == qn("w:p"):
                    break
                ancestor = ancestor.getparent()
            if ancestor is not p_el:
                continue
            if d.find(f"{{{NS['wp']}}}inline") is not None:
                inline_drawings.append(d)
            elif d.find(f"{{{NS['wp']}}}anchor") is not None:
                anchored_drawings.append(d)

        if not inline_drawings and not anchored_drawings:
            continue

        page = tracker.page_for(p_el)

        # Parse all anchored shapes for this paragraph.
        shapes_with_drawings: List[Tuple[FloatingShape, "object"]] = []
        for ad in anchored_drawings:
            # If the anchored drawing carries a picture (not a shape), the
            # existing "no floating images" rule applies. Detect this and
            # fail-fast with page number.
            gd = ad.find(f".//{{{NS['a']}}}graphicData")
            if gd is not None and gd.get("uri", "") == (
                "http://schemas.openxmlformats.org/drawingml/2006/picture"
            ):
                _fail_with_page(
                    f"{docx_path.name} contains a floating or anchored "
                    f"image. Only inline images are supported.",
                    page,
                )
            parsed = _parse_anchor_drawing(ad, page, docx_path)
            if parsed is not None:
                shapes_with_drawings.append((parsed, ad))

        # If this paragraph has anchored shapes but multiple inline images,
        # we can't unambiguously decide which image each shape overlays.
        if shapes_with_drawings and len(inline_drawings) > 1:
            _fail_with_page(
                f"{docx_path.name} has a paragraph with multiple inline "
                f"images plus floating shapes; v1 cannot disambiguate which "
                f"image each shape overlays.",
                page,
            )

        # Assign shapes to inline images by containment.
        per_image_overlays: Dict[int, List[FloatingShape]] = {
            id(d): [] for d in inline_drawings
        }
        for shape, ad in shapes_with_drawings:
            assigned = False
            for d in inline_drawings:
                extent = _inline_image_extent(d, page, docx_path)
                if _shape_within_extent(shape, extent):
                    per_image_overlays[id(d)].append(shape)
                    assigned = True
                    break
            if not assigned:
                _fail_with_page(
                    f"{docx_path.name} contains a floating shape "
                    f"({shape.kind}) that is not wholly inside an inline "
                    f"image. Only shapes contained within an inline image "
                    f"are supported.",
                    page,
                )

        for d in inline_drawings:
            yield (d, page, per_image_overlays[id(d)])


def _existing_assets_for_doc(git_dir: Path, md_name: str,
                             manifest_entry: Optional[Dict]) -> List[Path]:
    assets_dir = git_dir / ASSETS_DIR_NAME
    if not assets_dir.is_dir():
        return []
    known: List[Path] = []
    if manifest_entry:
        for img in manifest_entry.get("images", []):
            p = assets_dir / img["filename"]
            if p.exists():
                known.append(p)
    return known


def _allocate_image_filename(prefix: str, ext: str,
                             used: set, manifest_entry: Optional[Dict],
                             assets_dir: Path) -> str:
    """Pick a new image filename for the document, skipping any number
    currently in use either in the manifest or as a file on disk."""
    taken_numbers: set = set()
    # From the manifest (this document's previous images).
    if manifest_entry:
        for img in manifest_entry.get("images", []):
            m = re.match(rf"^{re.escape(prefix)}-image(\d+)\.[A-Za-z0-9]+$",
                         img["filename"])
            if m:
                taken_numbers.add(int(m.group(1)))
    # From disk (any file matching the prefix).
    if assets_dir.is_dir():
        for p in assets_dir.iterdir():
            m = re.match(
                rf"^{re.escape(prefix)}-image(\d+)\.[A-Za-z0-9]+$",
                p.name,
            )
            if m:
                taken_numbers.add(int(m.group(1)))
    # From the set we have already allocated in this run.
    for name in used:
        m = re.match(rf"^{re.escape(prefix)}-image(\d+)\.[A-Za-z0-9]+$", name)
        if m:
            taken_numbers.add(int(m.group(1)))
    n = 1
    while n in taken_numbers:
        n += 1
    return f"{prefix}-image{n}.{ext}"


def convert_docx_to_markdown(docx_path: Path, md_path: Path, git_dir: Path,
                             manifest: Dict) -> Dict[str, int]:
    """Convert a single .docx to .md, manage assets, update manifest.

    Returns counters: extracted, reused, removed.
    """
    counters = {"extracted": 0, "reused": 0, "removed": 0}

    # Open document.
    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        # python-docx raises a variety of errors; treat all as fatal.
        fail(
            f"Failed to read {docx_path.name}: {exc}. "
            f"If the file is open in Word, save and try again."
        )
        return counters  # unreachable

    tracker = PageTracker(doc)

    # Tracked changes: refuse to convert.
    _check_tracked_changes(doc, docx_path, tracker)

    # Extract and render all inline images in document order, compositing
    # any wholly-contained floating annotation shapes.
    extracted: List[ExtractedImage] = []
    drawing_order: List = []  # parallel list of inline-drawing elements
    for drawing_el, page, overlays in _collect_inline_images_with_overlays(
        doc, docx_path, tracker
    ):
        img = _extract_inline_image(
            drawing_el, doc, docx_path, page, overlays,
        )
        extracted.append(img)
        drawing_order.append(drawing_el)

    # Resolve filenames for each extracted image by hash, reusing existing
    # files where the content matches.
    md_name = md_path.name
    md_stem = md_path.stem
    manifest_entry = manifest["documents"].get(md_name)
    prev_by_hash: Dict[str, str] = {}
    prev_filenames: set = set()
    if manifest_entry:
        for entry in manifest_entry.get("images", []):
            prev_by_hash[entry["hash"]] = entry["filename"]
            prev_filenames.add(entry["filename"])

    assets_dir = git_dir / ASSETS_DIR_NAME

    final_filenames: List[str] = []
    final_alts: List[str] = []
    new_image_records: List[Dict] = []
    files_to_write: List[Tuple[Path, bytes]] = []
    used_in_this_run: set = set()

    for img in extracted:
        if img.hash_hex in prev_by_hash:
            # Reuse existing filename. Verify the file still exists; if it
            # was deleted out-of-band, re-create it.
            fname = prev_by_hash[img.hash_hex]
            target = assets_dir / fname
            if target.exists():
                counters["reused"] += 1
            else:
                files_to_write.append((target, img.rendered_bytes))
                counters["extracted"] += 1
        else:
            fname = _allocate_image_filename(
                md_stem, img.ext, used_in_this_run, manifest_entry, assets_dir
            )
            target = assets_dir / fname
            files_to_write.append((target, img.rendered_bytes))
            counters["extracted"] += 1

        used_in_this_run.add(fname)
        final_filenames.append(fname)
        final_alts.append(img.alt_text)
        new_image_records.append({
            "filename": fname,
            "hash": img.hash_hex,
            "width": img.width_px,
            "height": img.height_px,
        })

    # Render markdown.
    converter = DocxToMarkdown(doc, docx_path, final_filenames,
                               tracker=tracker)
    # Stash alts on the converter for emission.
    converter._image_alts = final_alts  # type: ignore[attr-defined]
    md_text = converter.render()

    # Validate that every referenced image will exist on disk after writes.
    referenced = set(final_filenames)
    referenced_on_disk = set()
    for fname in referenced:
        target = assets_dir / fname
        if target.exists() or any(p == target for p, _ in files_to_write):
            referenced_on_disk.add(fname)
    missing = referenced - referenced_on_disk
    if missing:
        fail(
            f"{docx_path.name}: markdown references images that would not "
            f"exist after write: {sorted(missing)}"
        )

    # Write images first (so the markdown points at real files).
    assets_dir.mkdir(parents=True, exist_ok=True)
    for target, blob in files_to_write:
        try:
            tmp = target.with_suffix(target.suffix + ".tmp")
            with open(tmp, "wb") as fp:
                fp.write(blob)
            os.replace(tmp, target)
        except OSError as exc:
            fail(f"Failed to write image {target}: {exc}")

    # Write markdown atomically.
    try:
        tmp_md = md_path.with_suffix(md_path.suffix + ".tmp")
        with open(tmp_md, "w", encoding="utf-8", newline="\n") as fp:
            fp.write(md_text)
        os.replace(tmp_md, md_path)
    except OSError as exc:
        # Clean up tmp on failure.
        try:
            if tmp_md.exists():
                tmp_md.unlink()
        except OSError:
            pass
        fail(f"Failed to write {md_path}: {exc}")

    # Now safe to remove stale image files for THIS document only.
    stale = prev_filenames - used_in_this_run
    for stale_name in stale:
        stale_path = assets_dir / stale_name
        if stale_path.exists():
            try:
                stale_path.unlink()
                counters["removed"] += 1
            except OSError as exc:
                fail(f"Failed to remove stale asset {stale_path}: {exc}")

    # Update manifest entry.
    manifest["documents"][md_name] = {
        "images": new_image_records,
        "updated": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
    }

    return counters


# ---------------------------------------------------------------------------
# Per-run orchestration
# ---------------------------------------------------------------------------

def _is_word_temp(filename: str) -> bool:
    return filename.startswith("~$")


def _list_docx_files(work_dir: Path) -> List[Path]:
    return sorted(
        p for p in work_dir.iterdir()
        if p.is_file()
        and p.suffix.lower() == ".docx"
        and not _is_word_temp(p.name)
    )


def _list_md_files(git_dir: Path) -> List[Path]:
    return sorted(
        p for p in git_dir.iterdir()
        if p.is_file() and p.suffix.lower() == ".md"
    )


def _build_case_insensitive_md_map(md_files: List[Path]) -> Dict[str, Path]:
    """Map stem.lower() -> md Path, failing on ambiguous duplicates."""
    result: Dict[str, Path] = {}
    seen: Dict[str, List[Path]] = {}
    for p in md_files:
        key = p.stem.lower()
        seen.setdefault(key, []).append(p)
    for key, ps in seen.items():
        if len(ps) > 1:
            names = ", ".join(p.name for p in ps)
            fail(
                f"Ambiguous .md files differing only in case: {names}. "
                f"Rename so stems are unique case-insensitively."
            )
        result[key] = ps[0]
    return result


def _check_docx_ambiguity(docx_files: List[Path]) -> None:
    seen: Dict[str, List[Path]] = {}
    for p in docx_files:
        seen.setdefault(p.stem.lower(), []).append(p)
    for key, ps in seen.items():
        if len(ps) > 1:
            names = ", ".join(p.name for p in ps)
            fail(
                f"Ambiguous .docx files differing only in case: {names}. "
                f"Rename so stems are unique case-insensitively."
            )


def run_togit(cfg: Dict[str, str]) -> None:
    git_dir = Path(cfg["GitDir"])
    work_dir = Path(cfg["WorkDir"])

    if not git_dir.is_dir():
        fail(f"GitDir does not exist or is not a directory: {git_dir}")
    if not work_dir.is_dir():
        fail(f"WorkDir does not exist or is not a directory: {work_dir}")

    docx_files = _list_docx_files(work_dir)
    md_files = _list_md_files(git_dir)

    _check_docx_ambiguity(docx_files)
    md_map = _build_case_insensitive_md_map(md_files)

    if not docx_files:
        print(f"No .docx files found in {work_dir}.")
        return

    manifest = load_manifest(git_dir)
    state = load_state(work_dir)

    totals = {
        "converted": 0,
        "unchanged": 0,
        "skipped": 0,
        "extracted": 0,
        "reused": 0,
        "removed": 0,
        "errors": 0,
    }

    for docx in docx_files:
        key = docx.stem.lower()
        md_match = md_map.get(key)
        if md_match is None:
            print(f"Skipping (no matching .md): {docx.name}")
            totals["skipped"] += 1
            continue

        if _docx_unchanged_since_last_sync(docx, state):
            print(f"Unchanged, skipping: {docx.name}")
            totals["unchanged"] += 1
            continue

        print(f"Converting: {docx.name} -> {md_match.name}")
        counters = convert_docx_to_markdown(docx, md_match, git_dir, manifest)
        totals["converted"] += 1
        totals["extracted"] += counters["extracted"]
        totals["reused"] += counters["reused"]
        totals["removed"] += counters["removed"]

        # Record this file's sync state and persist BOTH manifest and state
        # immediately. If a later file in the batch fails fast, the work
        # done up to that point is preserved on disk: re-running --togit
        # will skip already-synced files and retry the failed one.
        _record_sync(docx, state)
        save_manifest(git_dir, manifest)
        save_state(work_dir, state)

        print(f"Converted successfully: {md_match.name}")

    print()
    print(f"Converted:        {totals['converted']}")
    print(f"Unchanged:        {totals['unchanged']}")
    print(f"Skipped:          {totals['skipped']}")
    print(f"Images extracted: {totals['extracted']}")
    print(f"Images reused:    {totals['reused']}")
    print(f"Images removed:   {totals['removed']}")
    print(f"Errors:           {totals['errors']}")


# ---------------------------------------------------------------------------
# --fromgit operation
# ---------------------------------------------------------------------------

def run_fromgit(cfg: Dict[str, str]) -> None:
    git_dir = Path(cfg["GitDir"])
    work_dir = Path(cfg["WorkDir"])

    if not git_dir.is_dir():
        fail(f"GitDir does not exist or is not a directory: {git_dir}")
    if not work_dir.is_dir():
        fail(f"WorkDir does not exist or is not a directory: {work_dir}")

    # Root-level .md files only. Case-insensitive match for the extension on
    # Windows (which is case-insensitive anyway), but we accept any case here.
    md_files: List[Path] = sorted(
        p for p in git_dir.iterdir()
        if p.is_file() and p.suffix.lower() == ".md"
    )

    if not md_files:
        print(f"No .md files found in {git_dir}.")
        return

    created = 0
    skipped = 0

    for md in md_files:
        docx_name = md.stem + ".docx"
        docx_path = work_dir / docx_name
        if docx_path.exists():
            print(f"Skipping (already exists): {docx_path.name}")
            skipped += 1
            continue
        print(f"Converting: {md.name} -> {docx_name}")
        convert_markdown_to_docx(md, docx_path)
        created += 1

    print(f"\nDone. Created: {created}, Skipped: {skipped}.")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def parse_args(argv: List[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="mdtool",
        description=(
            "Manual markdown/Word working-copy utility. "
            "Operates on an already-synchronised local repository."
        ),
    )
    parser.add_argument(
        "--config",
        action="store_true",
        help="Enter interactive configuration mode.",
    )
    parser.add_argument(
        "--fromgit",
        action="store_true",
        help="Generate .docx working files from .md files in GitDir.",
    )
    parser.add_argument(
        "--togit",
        action="store_true",
        help="Update .md files in GitDir from edited .docx files in WorkDir.",
    )
    return parser.parse_args(argv)


def main(argv: Optional[List[str]] = None) -> None:
    args = parse_args(sys.argv[1:] if argv is None else argv)

    cfg_path = config_path()
    existing_cfg = load_config(cfg_path)

    # Configuration mode is triggered if config is missing OR --config flag.
    if args.config or existing_cfg is None:
        run_config_mode(existing_cfg)
        return  # run_config_mode calls sys.exit; safety net only.

    # Validate present config has both required keys.
    missing = [k for k in CONFIG_KEYS if k not in existing_cfg or not existing_cfg[k]]
    if missing:
        fail(
            f"Config {cfg_path} is missing required keys: {', '.join(missing)}. "
            f"Re-run with --config."
        )

    # Require an operational parameter when not in config mode.
    if args.fromgit and args.togit:
        fail("--fromgit and --togit are mutually exclusive.")
    if not args.fromgit and not args.togit:
        fail("No operation specified. Pass --fromgit, --togit, or --config.")

    if args.fromgit:
        run_fromgit(existing_cfg)
    else:
        run_togit(existing_cfg)


if __name__ == "__main__":
    main()